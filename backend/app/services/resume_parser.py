import logging
import re
from typing import Any
from typing import Dict
from typing import List

logger = logging.getLogger(__name__)


class ResumeParser:
    """Service for parsing resume data from uploaded files."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        self.phone_pattern = re.compile(r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})')
        self.url_pattern = re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')
        
    async def parse_resume(
        self, 
        file_content: bytes, 
        filename: str, 
        content_type: str
    ) -> Dict[str, Any]:
        """Parse resume from uploaded file."""
        try:
            file_extension = self._get_file_extension(filename)
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Extract text content based on file type
            text_content = await self._extract_text(file_content, file_extension, content_type)
            
            if not text_content or len(text_content.strip()) < 50:
                raise ValueError("Unable to extract meaningful content from file")
            
            # Parse structured data from text
            parsed_data = await self._parse_text_content(text_content)
            
            logger.info(f"Successfully parsed resume from {filename}")
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing resume from {filename}: {str(e)}")
            raise
    
    async def _extract_text(self, file_content: bytes, file_extension: str, content_type: str) -> str:
        """Extract text content from different file formats."""
        try:
            if file_extension == '.pdf':
                return await self._extract_from_pdf(file_content)
            elif file_extension in ['.docx', '.doc']:
                return await self._extract_from_word(file_content, file_extension)
            elif file_extension == '.txt':
                return file_content.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_extension} file: {str(e)}")
            raise
    
    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            from io import BytesIO

            import PyPDF2
            
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text_content = ""
            
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            return text_content
            
        except ImportError:
            # Fallback to pdfplumber if PyPDF2 not available
            try:
                from io import BytesIO

                import pdfplumber
                
                text_content = ""
                with pdfplumber.open(BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
                
                return text_content
                
            except ImportError:
                raise RuntimeError("No PDF parsing library available (PyPDF2 or pdfplumber)")
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise
    
    async def _extract_from_word(self, file_content: bytes, file_extension: str) -> str:
        """Extract text from Word document."""
        try:
            from io import BytesIO

            import docx
            
            if file_extension == '.docx':
                doc = docx.Document(BytesIO(file_content))
                text_content = ""
                
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                # Extract text from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + " "
                        text_content += "\n"
                
                return text_content
            else:
                # For .doc files, would need python-docx or other library
                raise ValueError("Legacy .doc format not supported. Please convert to .docx")
                
        except ImportError:
            raise RuntimeError("python-docx library not installed")
        except Exception as e:
            logger.error(f"Error extracting Word document text: {str(e)}")
            raise
    
    async def _parse_text_content(self, text: str) -> Dict[str, Any]:
        """Parse structured data from extracted text."""
        parsed_data = {
            "personal_info": {},
            "professional_summary": "",
            "experiences": [],
            "educations": [],
            "skills": [],
            "projects": [],
            "certifications": [],
            "languages": [],
            "references": []
        }
        
        try:
            # Extract personal information
            parsed_data["personal_info"] = self._extract_personal_info(text)
            
            # Extract professional summary/objective
            parsed_data["professional_summary"] = self._extract_summary(text)
            
            # Extract work experience
            parsed_data["experiences"] = self._extract_work_experience(text)
            
            # Extract education
            parsed_data["educations"] = self._extract_education(text)
            
            # Extract skills
            parsed_data["skills"] = self._extract_skills(text)
            
            # Extract projects
            parsed_data["projects"] = self._extract_projects(text)
            
            # Extract certifications
            parsed_data["certifications"] = self._extract_certifications(text)
            
            # Extract languages
            parsed_data["languages"] = self._extract_languages(text)
            
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing text content: {str(e)}")
            raise
    
    def _extract_personal_info(self, text: str) -> Dict[str, Any]:
        """Extract personal information from text."""
        info = {}
        
        # Extract email
        emails = self.email_pattern.findall(text)
        if emails:
            info["email"] = emails[0]
        
        # Extract phone
        phones = self.phone_pattern.findall(text)
        if phones:
            # Format phone number
            phone = phones[0]
            if isinstance(phone, tuple):
                info["phone"] = ''.join(phone)
            else:
                info["phone"] = phone
        
        # Extract URLs (LinkedIn, GitHub, websites)
        urls = self.url_pattern.findall(text)
        for url in urls:
            if 'linkedin.com' in url.lower():
                info["linkedin_url"] = url
            elif 'github.com' in url.lower():
                info["github_url"] = url
            elif not info.get("website") and not any(domain in url.lower() 
                                                   for domain in ['linkedin', 'github', 'facebook', 'twitter']):
                info["website"] = url
        
        # Extract name (this is tricky, using heuristics)
        lines = text.split('\n')[:10]  # Check first 10 lines
        for line in lines:
            line = line.strip()
            # Skip lines that are likely not names
            if (len(line.split()) in [2, 3] and 
                not any(char.isdigit() for char in line) and 
                '@' not in line and 
                len(line) > 5 and len(line) < 50):
                words = line.split()
                if all(word.replace('-', '').replace("'", '').isalpha() for word in words):
                    info["full_name"] = line
                    break
        
        return info
    
    def _extract_summary(self, text: str) -> str:
        """Extract professional summary or objective."""
        text.lower()
        
        # Common section headers for summaries
        summary_headers = [
            'professional summary', 'summary', 'profile', 'objective',
            'career objective', 'professional profile', 'about me'
        ]
        
        for header in summary_headers:
            pattern = rf'\b{re.escape(header)}\b\s*:?\s*(.*?)(?=\n\s*[A-Z][^a-z]*:|\n\s*\n|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                summary = match.group(1).strip()
                # Clean up the summary
                summary = re.sub(r'\n+', ' ', summary)
                summary = re.sub(r'\s+', ' ', summary)
                if len(summary) > 50:  # Ensure it's substantial
                    return summary[:500]  # Limit length
        
        return ""
    
    def _extract_work_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience from text."""
        experiences = []
        
        # This is a simplified extraction - real implementation would be more sophisticated
        exp_section = self._extract_section(text, ['experience', 'work experience', 'employment', 'career'])
        
        if exp_section:
            # Split by common patterns that indicate new job entries
            job_patterns = [
                r'\n(?=[A-Z][^a-z\n]*(?:at|@|\|)[^a-z\n]*[A-Z])',  # Title at Company
                r'\n(?=\d{4}\s*[-–]\s*\d{4})',  # Date ranges
                r'\n(?=\w+\s+\d{4})',  # Month Year
            ]
            
            # This would need more sophisticated parsing logic
            # For now, return a placeholder structure
            experiences.append({
                "company_name": "Company Name",
                "position_title": "Position Title", 
                "location": "",
                "start_date": None,
                "end_date": None,
                "is_current": False,
                "description": exp_section[:200] + "..." if len(exp_section) > 200 else exp_section,
                "achievements": [],
                "technologies": []
            })
        
        return experiences
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education from text."""
        educations = []
        
        edu_section = self._extract_section(text, ['education', 'academic background', 'qualifications'])
        
        if edu_section:
            # Look for degree patterns
            degree_patterns = [
                r'\b(?:Bachelor|Master|PhD|Ph\.D\.|Doctor|Associate|B\.A\.|B\.S\.|M\.A\.|M\.S\.|MBA)\b[^.]*',
                r'\b(?:BS|BA|MS|MA|PhD)\s+in\s+[^,\n]*'
            ]
            
            for pattern in degree_patterns:
                matches = re.findall(pattern, edu_section, re.IGNORECASE)
                for match in matches:
                    educations.append({
                        "institution_name": "University Name",  # Would extract from context
                        "degree": match.strip(),
                        "field_of_study": "",
                        "location": "",
                        "start_date": None,
                        "end_date": None,
                        "is_current": False,
                        "gpa": "",
                        "honors": "",
                        "description": "",
                        "courses": []
                    })
                    break  # For now, just take the first one
        
        return educations
    
    def _extract_skills(self, text: str) -> List[Dict[str, Any]]:
        """Extract skills from text."""
        skills = []
        
        skills_section = self._extract_section(text, [
            'skills', 'technical skills', 'core competencies', 
            'technologies', 'programming languages', 'tools'
        ])
        
        if skills_section:
            # Common programming languages and technologies
            tech_keywords = [
                'Python', 'JavaScript', 'Java', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'Rust',
                'React', 'Angular', 'Vue', 'Node.js', 'Django', 'Flask', 'Spring',
                'HTML', 'CSS', 'SQL', 'MongoDB', 'PostgreSQL', 'MySQL',
                'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Git'
            ]
            
            found_skills = []
            for skill in tech_keywords:
                if re.search(r'\b' + re.escape(skill) + r'\b', skills_section, re.IGNORECASE):
                    found_skills.append(skill)
            
            # Also extract comma-separated skills
            skill_lines = re.findall(r'[^\n]*[,•·]\s*[^\n]*', skills_section)
            for line in skill_lines:
                items = re.split(r'[,•·]\s*', line)
                for item in items:
                    item = item.strip()
                    if len(item) > 2 and len(item) < 30:
                        found_skills.append(item)
            
            # Convert to skill objects
            for skill_name in set(found_skills):  # Remove duplicates
                skills.append({
                    "name": skill_name,
                    "category": self._categorize_skill(skill_name),
                    "proficiency_level": None,
                    "proficiency_label": ""
                })
        
        return skills[:20]  # Limit to 20 skills
    
    def _extract_projects(self, text: str) -> List[Dict[str, Any]]:
        """Extract projects from text."""
        projects = []
        
        projects_section = self._extract_section(text, ['projects', 'personal projects', 'notable projects'])
        
        if projects_section:
            # Look for GitHub URLs as project indicators
            github_urls = re.findall(r'github\.com/[^\s]+', projects_section, re.IGNORECASE)
            
            for url in github_urls[:5]:  # Limit to 5 projects
                projects.append({
                    "name": "Project Name",  # Would extract from context
                    "description": "Project description",
                    "project_url": "",
                    "github_url": f"https://{url}",
                    "demo_url": "",
                    "start_date": None,
                    "end_date": None,
                    "is_ongoing": False,
                    "technologies": [],
                    "role": ""
                })
        
        return projects
    
    def _extract_certifications(self, text: str) -> List[Dict[str, Any]]:
        """Extract certifications from text."""
        certifications = []
        
        cert_section = self._extract_section(text, [
            'certifications', 'certificates', 'professional certifications',
            'licenses', 'credentials'
        ])
        
        if cert_section:
            # Common certification patterns
            cert_patterns = [
                r'\b(?:AWS|Azure|Google Cloud|GCP)\s+[^,\n]*(?:Certified|Certificate)',
                r'\b(?:PMP|CISSP|CISA|CCNA|MCSE|Oracle)\b[^,\n]*',
                r'\bCertified\s+[^,\n]*'
            ]
            
            for pattern in cert_patterns:
                matches = re.findall(pattern, cert_section, re.IGNORECASE)
                for match in matches[:5]:  # Limit to 5 certifications
                    certifications.append({
                        "name": match.strip(),
                        "issuing_organization": "Organization",  # Would extract from context
                        "credential_id": "",
                        "credential_url": "",
                        "issue_date": None,
                        "expiration_date": None,
                        "does_not_expire": True,
                        "description": ""
                    })
        
        return certifications
    
    def _extract_languages(self, text: str) -> List[Dict[str, Any]]:
        """Extract languages from text."""
        languages = []
        
        lang_section = self._extract_section(text, ['languages', 'language skills', 'spoken languages'])
        
        if lang_section:
            common_languages = [
                'English', 'Spanish', 'French', 'German', 'Italian', 'Portuguese',
                'Chinese', 'Japanese', 'Korean', 'Arabic', 'Russian', 'Hindi'
            ]
            
            proficiency_levels = {
                'native': 'Native',
                'fluent': 'Fluent', 
                'advanced': 'Advanced',
                'intermediate': 'Intermediate',
                'basic': 'Basic',
                'beginner': 'Beginner'
            }
            
            for language in common_languages:
                if re.search(r'\b' + re.escape(language) + r'\b', lang_section, re.IGNORECASE):
                    # Try to find proficiency level
                    proficiency = 'Fluent'  # Default
                    for level_key, level_value in proficiency_levels.items():
                        if re.search(rf'\b{language}\b.*?\b{level_key}\b|\b{level_key}\b.*?\b{language}\b', 
                                   lang_section, re.IGNORECASE):
                            proficiency = level_value
                            break
                    
                    languages.append({
                        "name": language,
                        "proficiency": proficiency
                    })
        
        return languages
    
    def _extract_section(self, text: str, headers: List[str]) -> str:
        """Extract a specific section from the resume text."""
        text.lower()
        
        for header in headers:
            # Look for section headers
            pattern = rf'\b{re.escape(header)}\b\s*:?\s*(.*?)(?=\n\s*[A-Z][^a-z]*:|\n\s*\n\s*[A-Z]|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                section_content = match.group(1).strip()
                if len(section_content) > 20:  # Ensure substantial content
                    return section_content
        
        return ""
    
    def _categorize_skill(self, skill_name: str) -> str:
        """Categorize a skill into a general category."""
        programming_languages = [
            'Python', 'JavaScript', 'Java', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'Rust', 'Swift', 'Kotlin'
        ]
        
        web_technologies = [
            'HTML', 'CSS', 'React', 'Angular', 'Vue', 'Node.js', 'Django', 'Flask', 'Spring'
        ]
        
        databases = ['SQL', 'MongoDB', 'PostgreSQL', 'MySQL', 'Oracle', 'Redis']
        
        cloud_platforms = ['AWS', 'Azure', 'GCP', 'Google Cloud', 'Heroku']
        
        tools = ['Git', 'Docker', 'Kubernetes', 'Jenkins', 'Jira', 'Confluence']
        
        if skill_name in programming_languages:
            return "Programming Languages"
        elif skill_name in web_technologies:
            return "Web Technologies"
        elif skill_name in databases:
            return "Databases"
        elif skill_name in cloud_platforms:
            return "Cloud Platforms"
        elif skill_name in tools:
            return "Tools & Technologies"
        else:
            return "Technical Skills"
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename."""
        return '.' + filename.split('.')[-1].lower() if '.' in filename else ''
    
    async def validate_resume_file(self, filename: str, file_size: int, content_type: str) -> bool:
        """Validate uploaded resume file."""
        # Check file extension
        file_extension = self._get_file_extension(filename)
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format. Supported formats: {', '.join(self.supported_formats)}")
        
        # Check file size (limit to 10MB)
        max_size = 10 * 1024 * 1024  # 10MB
        if file_size > max_size:
            raise ValueError(f"File size too large. Maximum size: {max_size / (1024*1024):.1f}MB")
        
        # Check content type
        allowed_content_types = {
            '.pdf': ['application/pdf'],
            '.docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
            '.doc': ['application/msword'],
            '.txt': ['text/plain']
        }
        
        if content_type not in allowed_content_types.get(file_extension, []):
            logger.warning(f"Content type mismatch: {content_type} for {file_extension}")
            # Don't raise error, just warn - some browsers send incorrect content types
        
        return True